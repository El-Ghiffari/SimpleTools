import pandas as pd
import docx
from docx import Document
from docx.shared import Pt
#from docx.oxml.ns import nsdecls
#from docx.oxml import parse_xml

def create_vulnerability_tables(excel_file, docx_file):
    # Read the Excel file into a pandas DataFrame
    df = pd.read_excel(excel_file)

    # Create a new Word document
    doc = Document()

    # Register Montserrat font in the document
    doc.styles['Normal'].font.name = 'Montserrat'

    # Iterate over the rows of the DataFrame
    for _, row in df.iterrows():
        # Get the vulnerability details from the row
        vulnerability = row['Vulnerability']
        severity = row['Severity']
        cvss = row['CVSS']
        impact = row['Impact']
        no = row['No']
        time = row['Time']
        ip_domain = row['IP/DOMAIN']
        endpoint = row['Endpoint']
        platform = row['Platform']
        description = row['Description']
        poc = row['PoC']
        recommendation = row['Recommendation']
        references = row['References']
        status = row['Status']

        # Add the vulnerability name as Heading 2 with Montserrat font
        doc.add_heading(vulnerability, level=2).style.font.name = 'Montserrat'

        # Create a table with two columns
        table = doc.add_table(rows=14, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        # Set the column widths
        # table.columns[0].width = Pt(100)
        # table.columns[1].width = Pt(400)

        # Set the cell values for each row
        table.cell(0, 0).paragraphs[0].add_run('No').bold = True
        table.cell(0, 1).text = str(no)

        table.cell(1, 0).paragraphs[0].add_run('Time').bold = True
        table.cell(1, 1).text = str(time)

        table.cell(2, 0).paragraphs[0].add_run('IP/DOMAIN').bold = True
        table.cell(2, 1).text = str(ip_domain)

        table.cell(3, 0).paragraphs[0].add_run('Endpoint').bold = True
        table.cell(3, 1).text = str(endpoint)

        table.cell(4, 0).paragraphs[0].add_run('Platform').bold = True
        table.cell(4, 1).text = str(platform)

        table.cell(5, 0).paragraphs[0].add_run('Vulnerability').bold = True
        table.cell(5, 1).text = str(vulnerability)

        table.cell(6, 0).paragraphs[0].add_run('Description').bold = True
        table.cell(6, 1).text = str(description)

        table.cell(7, 0).paragraphs[0].add_run('Severity').bold = True
        table.cell(7, 1).text = str(severity)

        table.cell(8, 0).paragraphs[0].add_run('CVSS').bold = True
        table.cell(8, 1).text = str(cvss)

        table.cell(9, 0).paragraphs[0].add_run('Impact').bold = True
        table.cell(9, 1).text = str(impact)

        table.cell(10, 0).paragraphs[0].add_run('PoC').bold = True
        poc_cell = table.cell(10, 1)
        poc_hyperlink = poc_cell.paragraphs[0].add_run()
        poc_hyperlink.text = str(poc)
        poc_hyperlink.font.name = 'Montserrat'
        poc_hyperlink.font.underline = True
        poc_hyperlink.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0xFF)
        poc_hyperlink.hyperlink = poc

        table.cell(11, 0).paragraphs[0].add_run('Recommendation').bold = True
        table.cell(11, 1).text = str(recommendation)

        table.cell(12, 0).paragraphs[0].add_run('References').bold = True
        references_cell = table.cell(12, 1)
        references_hyperlink = references_cell.paragraphs[0].add_run()
        references_hyperlink.text = str(references)
        references_hyperlink.font.name = 'Montserrat'
        references_hyperlink.font.underline = True
        references_hyperlink.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0xFF)
        references_hyperlink.hyperlink = references

        table.cell(13, 0).paragraphs[0].add_run('Status').bold = True
        table.cell(13, 1).text = str(status)

        # Apply Montserrat font to the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Montserrat'

    # Save the Word document
    doc.save(docx_file)

# Example usage
excel_file_path = input("Your path to excel file : ")
docx_file_path = 'report.docx'
create_vulnerability_tables(excel_file_path, docx_file_path)
